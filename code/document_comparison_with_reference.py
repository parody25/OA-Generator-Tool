import difflib
import openai
import streamlit as st
import tempfile
import fitz
from docx import Document
import os
import re
import json
import zipfile
import requests
import ollama
from openai import OpenAI
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageTemplate, Frame, PageBreak, PageTemplate, Spacer
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from datetime import datetime
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import Image

pdfmetrics.registerFont(TTFont('ArabicFont', "assets/ArabicFont/NotoKufiArabic-Regular.ttf"))

# Preprocessing text to standardize for comparison
def preprocess_text(text):
    return "\n".join([line.strip() for line in text.splitlines() if line.strip()])

# Extract text from Word documents, including tables
def extract_text_from_docx(doc):
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append("\t".join(row_text))
    return "\n".join(full_text)

# Extract text from PDFs, including structured content
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = []
    for page in doc:
        full_text.append(page.get_text("text"))
        for block in page.get_text("blocks") or []:
            if block[-1] == 0:
                full_text.append(block[4])
    return "\n".join(full_text)

# Generate HTML diff to display changes and common parts
def generate_html_diff(doc1_text, doc2_text):
    differ = difflib.HtmlDiff()
    html_diff = differ.make_file(doc1_text.splitlines(), doc2_text.splitlines(), context=True, numlines=1000)
    return html_diff

def process_points(content):
    styles = getSampleStyleSheet()
    style_body = styles["BodyText"]
    paragraphs = []

    if isinstance(content, str):
        # If content is a string, split it into numbered points
        points = re.split(r"(\d+\.\s)", content)
        buffer = ""
        for part in points:
            if re.match(r"^\d+\.\s$", part):
                if buffer:
                    paragraphs.append(Paragraph(buffer.strip(), style_body))
                buffer = f'<font color="red">{part.strip()}</font> '  # Highlight the number
            else:
                buffer += part
        if buffer:
            paragraphs.append(Paragraph(buffer.strip(), style_body))
    elif isinstance(content, list):
        # If content is a list, process each item as a point
        for item in content:
            if isinstance(item, str):
                # Add a red number if the item starts with a number
                if re.match(r"^\d+\.\s", item):
                    paragraphs.append(Paragraph(f'<font color="red">{item.strip()}</font>', style_body))
                else:
                    paragraphs.append(Paragraph(item.strip(), style_body))
    else:
        # If content is neither a string nor a list, convert it to a string
        paragraphs.append(Paragraph(str(content), style_body))

    return paragraphs

def process_points_arabic(content):
    from reportlab.lib.enums import TA_RIGHT
    styles = getSampleStyleSheet()
    style_body = ParagraphStyle(
        "RTLBody",
        parent=styles["BodyText"],
        alignment=TA_RIGHT,
        fontName="ArabicFont"
    )
    paragraphs = []

    if isinstance(content, str):
        # If content is a string, split it into numbered points
        points = re.split(r"(\d+\.\s)", content)
        buffer = ""
        for part in points:
            if re.match(r"^\d+\.\s$", part):
                if buffer:
                    paragraphs.append(Paragraph(buffer.strip(), style_body))
                buffer = f'<font color="red">{part.strip()}</font> '  # Highlight the number
            else:
                buffer = part + buffer  # Append to the beginning for RTL
        if buffer:
            paragraphs.append(Paragraph(buffer.strip(), style_body))
    elif isinstance(content, list):
        # If content is a list, process each item as a point
        for item in content:
            if isinstance(item, str):
                # Add a red number if the item starts with a number
                if re.match(r"^\d+\.\s", item):
                    paragraphs.append(Paragraph(f'<font color="red">{item.strip()}</font>', style_body))
                else:
                    paragraphs.append(Paragraph(item.strip(), style_body))
    else:
        # If content is neither a string nor a list, convert it to a string
        paragraphs.append(Paragraph(str(content), style_body))

    return paragraphs


def generate_templated_pdf(memo_content, template_path="constant/templatenew.json"):
    # Load the template JSON
    with open(template_path, "r") as f:
        template = json.load(f)
    
    # Normalize memo_content keys for consistent lookup
    memo_content = {
        key.lower().replace(" ", "_").replace(":", "").strip(): value
        for key, value in memo_content.items()
    }
    
    pdf_filename = "pdfs/templated_memo.pdf"
    left_margin = 50
    right_margin = 50
    bottom_margin = 50
    page_width, page_height = letter
    content_width = page_width - left_margin - right_margin

    # Define frames and templates
    first_page_frame = Frame(
        x1=left_margin,
        y1=bottom_margin,
        width=content_width,
        height=page_height - 320 - bottom_margin,
    )
    later_page_frame = Frame(
        x1=left_margin,
        y1=bottom_margin,
        width=content_width,
        height=page_height - 50 - bottom_margin,
    )
    first_page_template = PageTemplate(id="FirstPage", frames=[first_page_frame])
    later_pages_template = PageTemplate(id="LaterPages", frames=[later_page_frame])

    doc = SimpleDocTemplate(
        pdf_filename, 
        pagesize=letter,
        leftMargin=left_margin,
        rightMargin=right_margin,
        bottomMargin=bottom_margin,
    )
    doc.addPageTemplates([first_page_template, later_pages_template])
    
    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("Title", parent=styles["Heading1"], textColor=colors.red)
    style_section = ParagraphStyle("Section", parent=styles["Heading2"], textColor=colors.red)
    style_body = styles["BodyText"]

    elements = []
    print("Keys in memo_content:", memo_content.keys())
    # Iterate over each section in the template
    for section in template["template"]["sections"]:
        raw_content_key = section.get("content_key", "")
        title = section.get("title", "").strip()

        # Ensure content_key is always defined
        if raw_content_key:
            content_key = raw_content_key.lower().replace(" ", "_").replace(":", "").strip()
        else:
            content_key = None

        # Debugging Output
        print(f"Processing section: {title}")
        print(f"Original Content Key: {raw_content_key}, Normalized: {content_key}")

        # Lookup Content in memo_content (with safeguard)
        content = memo_content.get(content_key, "Content not available.") if content_key else "Content not available."
        print(f"Content Found: {content}")

        # Add section title and content to PDF
        elements.append(Paragraph(title, style_section))
        points = process_points(content)
        elements.extend(points)

        # Handle subsections if they exist
        if "subsections" in section:
            for subsection in section["subsections"]:
                raw_sub_content_key = subsection.get("content_key", "").strip()
                sub_title = subsection.get("title", "").strip()

                if raw_sub_content_key:
                    sub_content_key = raw_sub_content_key.lower().replace(" ", "_").replace(":", "").strip()
                    sub_content = memo_content.get(sub_content_key, "Content not available.")

                    print(f"Processing subsection: {sub_title}")
                    print(f"Original Sub Key: {raw_sub_content_key}, Normalized: {sub_content_key}")
                    print(f"Sub Content Found: {sub_content}")

                    elements.append(Paragraph(sub_title, style_section))
                    sub_points = process_points(sub_content)
                    elements.extend(sub_points)

    # Generate the PDF
    elements.append(PageBreak())
    doc.build(elements)

    # Add logo and footer after PDF generation
    doc = fitz.open(pdf_filename)
    w = 595
    h = 842
    footer_black_rect = fitz.Rect(w * 0.28, h * 0.9, w * 0.52, h)
    footer_yellow_rect = fitz.Rect(w * 0.52, h * 0.9 , w * 0.68, h)
    logo_rect = fitz.Rect(w * 0.7, h * 0.03, w * 0.9, h * 0.09)
    official_use_rect = fitz.Rect(w * 0.35, h * 0.02, w * 0.65, h * 0.06)
    bank_muscat_text_rect = fitz.Rect(w * 0.1, h * 0.1, w * 0.4, h * 0.15)
    operations_advice_rect = fitz.Rect(w * 0.7, h * 0.1, w * 0.9, h * 0.15)
    horizontal_line_rect = fitz.Rect(w * 0.1, h * 0.16, w * 0.9, h * 0.161)
    date_rect = fitz.Rect(w * 0.1, h * 0.17, w * 0.5, h * 0.22)
    to_rect = fitz.Rect(w * 0.1, h * 0.22, w * 0.9, h * 0.26)
    cc_rect = fitz.Rect(w * 0.1, h * 0.26, w * 0.9, h * 0.30)
    subject_rect = fitz.Rect(w * 0.1, h * 0.30, w * 0.9, h * 0.34)
    horizontal_line2_rect = fitz.Rect(w * 0.1, h * 0.35, w * 0.9, h * 0.355)
    logo_path = "assets/bm_logo.png"
    current_date = datetime.now()
    formatted_date = current_date.strftime("%B %d, %Y")

    # Text content
    footer_text_black = "This Document is classified as"
    footer_text_yellow = "Official Use"
    official_use_text = "Official Use"
    bank_muscat_text = "Bank Muscat"
    operations_advice_text = "Operations Advice\nNo. 001/2025"
    date_text = f"Date: {formatted_date}"
    to_text = "To: All Branches / Account Opening Team / Compliance Team / Other Teams"
    cc_text = "Cc: Chief Internal Audit / Regional Managers / Head Digital Products / Others"
    subject_text = "Subject: Launch of Business Account Origination using IBPS System"

    for page_num, page in enumerate(doc):
        page.insert_image(logo_rect, filename=logo_path)
        page.insert_textbox(official_use_rect, official_use_text, fontsize=12, fontname="Helvetica", color=(1, 1, 0), align=1)
        page.insert_textbox(footer_black_rect, footer_text_black, fontsize=10, fontname="Helvetica", color=(0, 0, 0), align=1)
        page.insert_textbox(footer_yellow_rect, footer_text_yellow, fontsize=10, fontname="Helvetica-Bold", color=(1, 1, 0), align=1)

        if page_num == 0:
            page.insert_textbox(bank_muscat_text_rect, bank_muscat_text, fontsize=14, fontname="Helvetica-Bold", color=(0, 0, 0), align=0)
            page.insert_textbox(operations_advice_rect, operations_advice_text, fontsize=12, fontname="Helvetica", color=(0, 0, 0), align=2)
            page.draw_line((horizontal_line_rect.x0, horizontal_line_rect.y0), (horizontal_line_rect.x1, horizontal_line_rect.y0), color=(0, 0, 0), width=3.0)
            page.insert_textbox(date_rect, date_text, fontsize=10, fontname="Helvetica", color=(0, 0, 0), align=0)
            page.insert_textbox(to_rect, to_text, fontsize=10, fontname="Helvetica", color=(0, 0, 0), align=0)
            page.insert_textbox(cc_rect, cc_text, fontsize=10, fontname="Helvetica", color=(0, 0, 0), align=0)
            page.insert_textbox(subject_rect, subject_text, fontsize=10, fontname="Helvetica", color=(0, 0, 0), align=0)
            page.draw_line((horizontal_line2_rect.x0, horizontal_line2_rect.y0), (horizontal_line2_rect.x1, horizontal_line2_rect.y0), color=(0, 0, 0), width=3.0)

    doc.save("pdfs/templated_memo_final.pdf")
    print("PDF generation completed successfully.")

def generate_arabic_pdf(memo_content, template_path="constant/templatenew.json"):
    # Load the template JSON
    pdfmetrics.registerFont(TTFont('ArabicFont', "assets/ArabicFont/NotoKufiArabic-Regular.ttf"))
    with open(template_path, "r") as f:
        template = json.load(f)

    pdf_filename = "pdfs/arabic_memo.pdf"
    left_margin = 50
    right_margin = 50
    bottom_margin = 50
    page_width, page_height = letter
    content_width = page_width - left_margin - right_margin

    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("Title", parent=styles["Heading1"], fontName="ArabicFont", textColor=colors.red, alignment=2)
    style_section = ParagraphStyle("Section", parent=styles["Heading2"], fontName="ArabicFont", textColor=colors.red, alignment=2)
    style_body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="ArabicFont", alignment=2)

    elements = []

    for section in template["template"]["sections"]:
        content_key = section.get("content_key", None)

        if content_key:
            content = memo_content.get(content_key, "المحتوى غير متوفر.")
            translated_content = translate_to_arabic(content)
            translated_content = translated_content.encode("utf-8").decode("utf-8")
            if section["title"] != "Summary of Changes":
                translated_title = translate_to_arabic(section["title"])
                elements.append(Paragraph(translated_title, style_section))
            points = process_points_arabic(translated_content)
            elements.extend(points)
            #elements.append(Paragraph(translated_content, style_body))

        if "subsections" in section:
            for subsection in section["subsections"]:
                sub_content_key = subsection.get("content_key", None)

                if sub_content_key:
                    sub_content = memo_content.get(sub_content_key, "المحتوى غير متوفر.")
                    translated_sub_content = translate_to_arabic(sub_content)
                    translated_subtitle = translate_to_arabic(subsection["title"])
                    elements.append(Paragraph(translated_subtitle, style_section))
                    points = process_points_arabic(translated_content)
                    elements.extend(points)
                    #elements.append(Paragraph(translated_sub_content, style_body))
                else:
                    translated_subtitle = translate_to_arabic(subsection["title"])
                    elements.append(Paragraph(translated_subtitle, style_section))
                    elements.append(Paragraph("المحتوى غير متوفر.", style_body))

    elements.append(PageBreak())
    # Header/Footer function
    def add_header_footer(canvas, doc):
        canvas.saveState()

        # Header content (common to all pages)
        canvas.setFont("ArabicFont", 12)
        canvas.setFillColor(colors.yellow)
        canvas.drawString(page_width / 2 - 50, page_height - 30, "للاستخدام الرسمي")
        canvas.setFillColor(colors.black)
        canvas.drawImage("assets/bm_logo.png", page_width - 195, page_height - 65, width=156, height=39)  # Top right

        # Footer content (common to all pages)
        canvas.setFont("ArabicFont", 10)
        canvas.drawString(page_width / 2 - 100, 30, "هذا المستند للاستخدام الرسمي فقط")
        
        if doc.page == 1:
            # Content for the first page only
            canvas.drawString(page_width - 250, page_height - 80, "مشورة العمليات رقم 001/2025")
            canvas.drawString(left_margin, page_height - 80, "بنك مسقط")
            canvas.setStrokeColor(colors.black)
            canvas.setLineWidth(3)
            canvas.line(left_margin, page_height - 90, page_width - right_margin, page_height - 90)
            # Section titles (first page only)
            canvas.drawRightString(page_width - right_margin, page_height - 110, "التاريخ: " + datetime.now().strftime('%Y-%m-%d'))
            canvas.drawRightString(page_width - right_margin, page_height - 135, "إلى: جميع الفروع / فريق فتح الحسابات / فريق الامتثال / الفرق الأخرى")
            canvas.drawRightString(page_width - right_margin, page_height - 160, "نسخة إلى: المدقق الداخلي الرئيسي / المديرون الإقليميون / رئيس المنتجات الرقمية / آخرون")
            canvas.drawRightString(page_width - right_margin, page_height - 185, "الموضوع: إطلاق حساب الأعمال باستخدام نظام IBPS")
            canvas.line(left_margin, page_height - 195, page_width - right_margin, page_height - 195)
        canvas.restoreState()

    # Define frames and templates
    first_page_frame = Frame(
        x1=left_margin,
        y1=bottom_margin,
        width=content_width,
        height=page_height - 225 - bottom_margin,
        rightPadding=0, leftPadding=0
    )
    later_page_frame = Frame(
        x1=left_margin,
        y1=bottom_margin,
        width=content_width,
        height=page_height - 70 - bottom_margin,
        rightPadding=0, leftPadding=0
    )
    first_page_template = PageTemplate(id="FirstPage", frames=[first_page_frame], onPage=add_header_footer)
    later_pages_template = PageTemplate(id="LaterPages", frames=[later_page_frame], onPage=add_header_footer)

    # Build the PDF
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        leftMargin=left_margin,
        rightMargin=right_margin,
        bottomMargin=bottom_margin,
    )
    doc.addPageTemplates([first_page_template, later_pages_template])
    doc.build(elements)

    print("Arabic PDF generation completed successfully.")

def generate_gpt4_memo(doc1_text, doc2_text, responsibilties_selected, stakeholders_context, template_path="constant/templatenew.json"):
    # Load the template JSON
    with open(template_path, "r") as f:
        template = json.load(f)
    
    differ = difflib.Differ()
    diff = list(differ.compare(doc1_text.splitlines(), doc2_text.splitlines()))
    
    additions = [line[2:] for line in diff if line.startswith("+ ")]
    deletions = [line[2:] for line in diff if line.startswith("- ")]

    responsibilities_string = ", ".join([f"{s} ({stakeholders_context[s]})" for s in responsibilties_selected])
    keys_string = ", ".join([f'"{s.lower().replace(" ", "_")}_responsibilities"' for s in responsibilties_selected])

    # Dynamically include sections from the template in the prompt
    sections_prompt = ""
    for section in template["template"]["sections"]:
        title = section.get("title", "")
        content_key = section.get("content_key", "")
        sections_prompt += f"- **{title}**: Provide content for {content_key}.\n"

    prompt = f"""
    You are an AI assistant generating structured operational advice based on document changes.

    **Document Differences:**
    Additions: {', '.join(additions)}
    Deletions: {', '.join(deletions)}

    **Required Sections:**
    {sections_prompt}
    - **Responsibilities** ({responsibilities_string})

    **Response Format:** JSON with the following keys:
    {keys_string}.
    Provide the response in a **single string** formatted as **point-by-point**, with each point starting on a new line and prefixed with a number **example 1,2,3 etc**.
    """

    print("Prompt", prompt)
    try:
        openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        completion = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant for generating structured operational advice memos based on document changes."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=10000,
            temperature=0.7,
        )
        content = completion.choices[0].message.content
        print("Content", content)
        clean_content = content.strip().replace("```json", "").replace("```", "").strip()
        try:
            memo_content = json.loads(clean_content)
            # Debug: Print the generated memo_content
            print("Generated memo_content:", memo_content)
            return memo_content
        except json.JSONDecodeError:
            # Attempt to extract JSON fragment
            json_fragment = extract_json(clean_content)
            if json_fragment:
                return json.loads(json_fragment)
            else:
                raise ValueError("Response does not contain valid JSON.")
    except Exception as e:
        st.error(f"Error generating memo: {e}")
        raise

def generate_deepseek_memo(doc1_text, doc2_text, responsibilties_selected, stakeholders_context):
    differ = difflib.Differ()
    diff = list(differ.compare(doc1_text.splitlines(), doc2_text.splitlines()))
    
    additions = [line[2:] for line in diff if line.startswith("+ ")]
    deletions = [line[2:] for line in diff if line.startswith("- ")]

    responsibilities_string = ""   # To add in the **Responsibilities** line in prompt
    keys_string = ""   # To add in the JSON format keys

    for s in responsibilties_selected:
        responsibilities_string += f"{s} ({stakeholders_context[s]}), "
        keys_string += f'"{s.lower().replace(" ", "_")}_responsibilities", '

    # Remove the last ", "
    responsibilities_string = responsibilities_string.rstrip(", ")
    keys_string = keys_string.rstrip(", ")
    
    prompt = f"""
    You are an assistant helping with document comparison. Below are the differences between two documents:

    Additions:
    {', '.join(additions)}

    Deletions:
    {', '.join(deletions)}

    Based on these differences, provide detailed outputs for the following sections. Ensure the outputs are specific, actionable, and tailored to the context of operational advice:
    - **Summary of the changes**: Provide a well-written response summarizing the key changes between the documents.
    - **Highlights of the new service**: Focus on specific features or improvements introduced by the changes. Highlight their significance or utility.
    - **Potential impact or concerns**: Identify specific operational, compliance, or customer-related risks or benefits caused by these changes. Provide examples or scenarios where applicable.
    - **Actionable recommendations**: Provide actions that teams should take to address the changes, mitigate risks, or maximize benefits.
    - **Responsibilities** ({responsibilities_string}): Specify concrete tasks or duties for each team, ensuring these are directly tied to their roles and the changes.

    Provide the response in **strict JSON format** with the following keys:
    "summary_of_changes", "highlights", "impact", "recommendations", {keys_string}.

    - The value for "summary_of_changes" should be a paragraph.
    - The values for "highlights", "impact", "recommendations", and {keys_string} should be **single strings** formatted as point-by-point lists, with each point starting on a new line and prefixed with a number (e.g., 1, 2, 3, etc.).

    **Important**: 
    - Return only valid JSON.
    - Do not include any additional text, explanations, or markdown formatting.
    - Ensure all values are strings, even for sections like "highlights", "impact", and "recommendations". Combine multiple points into a single string with each point on a new line.
    """

    print("Prompt:", prompt)

    try:
        response = ollama.generate(
            model="deepseek-r1:7b",
            prompt=prompt,
            system="You are a helpful assistant for generating structured operational advice memos based on document changes.",
            options={
                "max_tokens": 12000,
                "temperature": 0.7,
            }
        )

        # Extract response content
        content = response['response'].strip()
        print("Content", content)
        # Attempt to extract and parse JSON
        json_response = extract_json_from_response(content)
        if json_response:
            print("Json Response", json_response)
            return json_response
        else:
            print("Failed to extract valid JSON from the Deepseek's response.")
            return None

    except Exception as e:
        st.error(f"Error generating memo: {e}")
        raise

def extract_json_from_response(response_text):
    # Use regex to find a JSON object in the response
    json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
    if json_match:
        try:
            # Attempt to parse the JSON
            return json.loads(json_match.group(0))
        except json.JSONDecodeError:
            # If parsing fails, return None
            return None
    return None

def extract_json(text):
    json_pattern = r'({.*?}|\[.*?\])'  # Match both object `{}` and array `[]`
    matches = re.findall(json_pattern, text, re.DOTALL)
    if matches:
        for match in matches:
            try:
                json.loads(match)  # Validate JSON
                return match
            except json.JSONDecodeError:
                continue
    return None

def update_template_with_reference(reference_structure, template_path="constant/templatenew.json"):
    # Load the existing template
    with open(template_path, "r") as f:
        template = json.load(f)
    
    # Update the template with the reference structure
    for section in reference_structure["template"]["sections"]:
        # Check if the section already exists in the template
        existing_section = next((s for s in template["template"]["sections"] if s["title"] == section["title"]), None)
        if not existing_section:
            # Add the new section to the template
            template["template"]["sections"].append(section)
    
    # Save the updated template
    with open(template_path, "w") as f:
        json.dump(template, f, indent=4)

def extract_structure_from_reference(reference_doc):
    structure = {"template": {"sections": []}}
    
    if reference_doc.type == "application/pdf":
        doc = fitz.open(stream=reference_doc.read(), filetype="pdf")
        for page in doc:
            text = page.get_text("text")
            # Extract headings and content (this is a simple example)
            headings = re.findall(r"^\s*(.*?)\s*$", text, re.MULTILINE)
            for heading in headings:
                structure["template"]["sections"].append({
                    "title": heading,
                    "content_key": heading.lower().replace(" ", "_")
                })
    
    elif reference_doc.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        doc = Document(reference_doc)
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                structure["template"]["sections"].append({
                    "title": para.text,
                    "content_key": para.text.lower().replace(" ", "_")
                })
    
    return structure

def translate_to_arabic(text):
    openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    def translate_single_item(item):
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are a helpful translator. Translate the following text to Arabic."},
                    {"role": "user", "content": str(item)}
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            print(f"Error during translation: {e}")
            return item  # Return the original item if translation fails

    # Handle different input types
    if isinstance(text, str):
        return translate_single_item(text)
    elif isinstance(text, list):
        return [translate_single_item(item) for item in text]
    elif isinstance(text, dict):
        return {key: translate_single_item(value) for key, value in text.items()}
    else:
        print(f"Unsupported input type: {type(text)}")
        return text

# Main Streamlit application logic
def document_comparison_with_reference():
    st.header("Operational Advice Generator With Reference DOC 📝")
    # File uploader for reference document
    reference_doc = st.file_uploader("Upload Reference Document (PDF, DOCX)", type=["pdf", "docx"])

    # File uploader
    documents = st.file_uploader("Upload documents (PDF, DOCX, TXT, PPTX)", 
                                  type=["pdf", "docx", "txt", "pptx"], 
                                  accept_multiple_files=True)

    if not documents:
        st.info("Please upload at least two documents.")
        return
    if reference_doc:
        # Extract structure from the reference document
        reference_structure = extract_structure_from_reference(reference_doc)
        # Update the template with the reference structure
        update_template_with_reference(reference_structure)
        st.success("Reference document processed and template updated!")

    all_texts = []
    doc_names = []
    for doc in documents:
        try:
            if doc.type == "application/pdf":
                with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                    tmp_file.write(doc.getvalue())
                    tmp_file_path = tmp_file.name
                all_texts.append(preprocess_text(extract_text_from_pdf(tmp_file_path)))
                doc_names.append(doc.name)

            elif doc.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                docx = Document(doc)
                all_texts.append(preprocess_text(extract_text_from_docx(docx)))
                doc_names.append(doc.name)

            elif doc.type == "text/plain":
                all_texts.append(preprocess_text(doc.read().decode("utf-8")))
                doc_names.append(doc.name)

            else:
                st.warning(f"Unsupported file type: {doc.type}")
        except Exception as e:
            st.warning(f"Error processing {doc.name}: {e}")

    if len(all_texts) < 2:
        st.warning("Please upload at least two valid documents.")
        return

    st.subheader("Stakeholders Involved")
    with open("constant/stakeholders.json", "r") as config_file:
        config = json.load(config_file)
    #responsibilties = ["Branch", "Service Manager", "Account Opening Team", "Compliance Team", "Online Banking Team"]
    stakeholders = config.get("stakeholders", {})
    responsibilties = list(stakeholders.keys())
    stakeholders_context = stakeholders
    responsibilties_selected = []
    # cols = st.columns(len(responsibilties))
    with st.form(key="responsibilties_form"):
        for responsibility in responsibilties:
            is_selected = st.checkbox(responsibility)
            if is_selected:
                responsibilties_selected.append(responsibility)
        done_button = st.form_submit_button(label="Done")
    model_choice = st.radio("Choose a model to generate the memo:", ["GPT-4o", "DeepSeek"], index=0)
    # Select documents for comparison
    if done_button:
        selected_docs = st.multiselect("Select two documents for comparison:", options=doc_names, default=doc_names[:2])
        if len(selected_docs) == 2:
            doc1_idx, doc2_idx = map(doc_names.index, selected_docs)
            doc1_text, doc2_text = all_texts[doc1_idx], all_texts[doc2_idx]

            # Perform comparison and generate HTML diff
            st.subheader("Comparison Report")
            with st.spinner("Comparing documents..."):
                html_diff = generate_html_diff(doc1_text, doc2_text)
                st.markdown("### Difference Viewer")
                st.components.v1.html(html_diff, height=600, scrolling=True)
        

            # Generate Operational Advice Memo using GPT-4
            with st.spinner("Generating memo content..."):
                try:
                    if model_choice == "GPT-4o":
                        memo_content = generate_gpt4_memo(doc1_text, doc2_text, responsibilties_selected, stakeholders_context)
                    else:
                        memo_content = generate_deepseek_memo(doc1_text, doc2_text, responsibilties_selected, stakeholders_context)
                    
                    st.success("Memo content generated successfully!")
                except Exception as e:
                    st.error(f"Failed to generate memo content: {e}")
                    return

            # Generate and provide PDF memo
            with st.spinner("Generating PDF memo..."):
                try:
                    english_pdf_path = "pdfs/templated_memo_final.pdf"
                    arabic_pdf_path = "pdfs/arabic_memo.pdf"

                    # Generate the PDFs using your custom function
                    generate_templated_pdf(memo_content)
                    generate_arabic_pdf(memo_content)
                    #generate_templated_pdf(memo_content)  # Pass memo_content here
                    st.success("Memo PDF generated successfully!")
                except Exception as e:
                    st.error(f"Failed to generate PDF memo: {e}")
                    return
            # Create a ZIP file containing both PDFs
            zip_file_path = "pdfs/memo_documents.zip"
            with zipfile.ZipFile(zip_file_path, "w") as zipf:
                zipf.write(english_pdf_path, arcname="operational_advice_english.pdf")
                zipf.write(arabic_pdf_path, arcname="operational_advice_arabic.pdf")

            # Provide download link for the generated PDF
            with open(zip_file_path, "rb") as zip_file:
                st.download_button(
                    label="Download Memo PDFs (English & Arabic)",
                    data=zip_file,
                    file_name="memo_documents.zip",
                    mime="application/zip"
                )
            #with open("pdfs/templated_memo_final.pdf", "rb") as pdf_file:
                #st.download_button("Download Memo PDF", pdf_file, file_name="operational_advice_memo.pdf")

        else:
            st.info("Select exactly two documents for comparison.")

if __name__ == "__main__":
    document_comparison_with_reference()
