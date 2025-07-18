import openai
import streamlit as st
import tempfile
import fitz
from docx import Document
import os
import json
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageTemplate, Frame, PageBreak
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from datetime import datetime
from openai import OpenAI

# Preprocess text to clean and standardize
def preprocess_text(text):
    return "\n".join([line.strip() for line in text.splitlines() if line.strip()])

# Extract text from Word documents
def extract_text_from_docx(doc):
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            full_text.append("\t".join(row_text))
    return "\n".join(full_text)

# Extract text from PDFs
def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    full_text = []
    for page in doc:
        full_text.append(page.get_text("text"))
    return "\n".join(full_text)

# Call GPT-4 API to generate a procedure manual
def generate_procedure_manual(content):
    prompt = f"""
    You are an expert assistant. Based on the following operations advice document, generate a detailed procedure manual. 
    Structure the manual with headings, subheadings, and step-by-step instructions. Ensure it's professional and comprehensive.

    Operations Advice Content:
    {content}
    """
    try:
        openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        completion = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a helpful assistant for generating procedure manuals."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=5000,
            temperature=0.7
        )
        content = completion.choices[0].message.content
        return content
    except Exception as e:
        st.error(f"Error generating procedure manual: {e}")
        raise

# Generate a PDF from the procedure manual
def generate_pdf(procedure_content, template_path="constant/procedure_manual_template.json"):
    # Load the template JSON
    with open(template_path, "r") as f:
        template = json.load(f)
    
    # File name and layout setup
    pdf_filename = "pdfs/templated_procedure_manual.pdf"
    left_margin = 50
    right_margin = 50
    bottom_margin = 50
    page_width, page_height = letter
    content_width = page_width - left_margin - right_margin
    first_page_frame = Frame(
        x1=left_margin,
        y1=bottom_margin,
        width=content_width,
        height=page_height - 150 - bottom_margin,
    )
    later_page_frame = Frame(
        x1=left_margin,
        y1=bottom_margin,
        width=content_width,
        height=page_height - 50 - bottom_margin,
    )
    first_page_template = PageTemplate(id="FirstPage", frames=[first_page_frame])
    later_pages_template = PageTemplate(id="LaterPages", frames=[later_page_frame])

    # Document settings
    doc = SimpleDocTemplate(
        pdf_filename,
        pagesize=letter,
        leftMargin=left_margin,
        rightMargin=right_margin,
        bottomMargin=bottom_margin,
    )
    doc.addPageTemplates([first_page_template, later_pages_template])
    
    # Styles
    styles = getSampleStyleSheet()
    style_title = ParagraphStyle("Title", parent=styles["Heading1"], textColor=colors.red)
    style_section = ParagraphStyle("Section", parent=styles["Heading2"], textColor=colors.darkblue)
    style_body = styles["BodyText"]

    # Add content based on the template
    elements = [Paragraph("Procedure Manual", style_title)]
    
    for section in template["template"]["sections"]:
        # Section Title
        elements.append(Paragraph(section["title"], style_section))
        
        # Section Content
        content_key = section.get("content_key", None)
        if content_key:
            content = procedure_content.get(content_key, "Content not available.")
            elements.extend(process_points(content))  # Format as points if applicable
        
        # Subsections
        if "subsections" in section:
            for subsection in section["subsections"]:
                elements.append(Paragraph(subsection["title"], style_section))
                sub_content_key = subsection.get("content_key", None)
                if sub_content_key:
                    sub_content = procedure_content.get(sub_content_key, "Content not available.")
                    elements.extend(process_points(sub_content))  # Format as points if applicable

    elements.append(PageBreak())
    doc.build(elements)
        # Add logo and footer after PDF generation
    doc = fitz.open(pdf_filename)
    w = 595
    h = 842
    footer_black_rect = fitz.Rect(w * 0.28, h * 0.9, w * 0.52, h)
    footer_yellow_rect = fitz.Rect(w * 0.52, h * 0.9 , w * 0.68, h)
    logo_rect = fitz.Rect(w * 0.7, h * 0.03, w * 0.9, h * 0.09)
    official_use_rect = fitz.Rect(w * 0.35, h * 0.02, w * 0.65, h * 0.06)
    branch_operation_produre_rect = fitz.Rect(w * 0.1, h * 0.1, w * 0.4, h * 0.15)
    #date_rect = fitz.Rect(w * 0.1, h * 0.17, w * 0.5, h * 0.22)
    logo_path = "assets/bm_logo.png"
    current_date = datetime.now()
    formatted_date = current_date.strftime("%B %d, %Y")
    # Text content
    footer_text_black = "This Document is classified as"
    footer_text_yellow = "Official Use"
    official_use_text = "Official Use"
    branch_operation_produre_text = "BRANCH OPERATION PROCEDURE MANUAL"
    #date_text = f"Date: {formatted_date}"

    for page_num, page in enumerate(doc):
        page.insert_image(logo_rect, filename=logo_path)
        #page.insert_textbox(footer_rect, "Confidential - Bank Muscat", fontsize=10, align=fitz.TEXT_ALIGN_CENTER)
        page.insert_textbox(official_use_rect, official_use_text, fontsize=12, fontname="Helvetica", color=(1, 1, 0), align=1)
        page.insert_textbox(footer_black_rect, footer_text_black, fontsize=10, fontname="Helvetica", color=(0, 0, 0), align=1)
        page.insert_textbox(footer_yellow_rect, footer_text_yellow, fontsize=10, fontname="Helvetica-Bold", color=(1, 1, 0), align=1)

        if page_num == 0:
            page.insert_textbox(branch_operation_produre_rect, branch_operation_produre_text, fontsize=14, fontname="Helvetica", color=(0, 0, 0), align=0)

    doc.save("pdfs/templated_procedure_manual_final.pdf")
    print("PDF generation completed successfully.")
    return pdf_filename

def process_points(content):
    """Process content into numbered points or paragraphs."""
    styles = getSampleStyleSheet()
    style_body = styles["BodyText"]
    points = re.split(r"(\d+\.\s)", content)

    paragraphs = []
    buffer = ""
    for part in points:
        if re.match(r"^\d+\.\s$", part):
            if buffer:
                paragraphs.append(Paragraph(buffer.strip(), style_body))
            buffer = f"<b>{part.strip()}</b> "
        else:
            buffer += part

    if buffer:
        paragraphs.append(Paragraph(buffer.strip(), style_body))

    return paragraphs

def generate_procedure_manual_page():
    st.header("Generate Procedure Manual from Operations Advice 📄")

    # File uploader for multiple documents
    documents = st.file_uploader(
        "Upload documents (PDF, DOCX, TXT, PPTX)",
        type=["pdf", "docx", "txt", "pptx"],
        accept_multiple_files=True
    )

    if documents:
        try:
            all_extracted_text = []

            # Extract text from each uploaded document
            for document in documents:
                if document.type == "application/pdf":
                    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                        tmp_file.write(document.getvalue())
                        tmp_file_path = tmp_file.name
                    all_extracted_text.append(preprocess_text(extract_text_from_pdf(tmp_file_path)))

                elif document.type in [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/msword",
                ]:
                    docx = Document(document)
                    all_extracted_text.append(preprocess_text(extract_text_from_docx(docx)))

                elif document.type == "text/plain":
                    all_extracted_text.append(preprocess_text(document.read().decode("utf-8")))

                else:
                    st.warning(f"Unsupported file type: {document.type}")
                    continue

            if not all_extracted_text:
                st.warning("No valid content could be extracted from the uploaded documents.")
                return

            # Combine text from all documents
            combined_text = "\n\n".join(all_extracted_text)

            # Generate procedure manual using GPT-4
            st.subheader("Generating Procedure Manual...")
            with st.spinner("Processing the documents and generating the manual..."):
                procedure_manual = generate_procedure_manual(combined_text)

            st.success("Procedure manual generated successfully!")

            # Display the manual in the app
            st.subheader("Generated Procedure Manual:")
            st.write(procedure_manual)

            # Generate and provide the PDF
            st.subheader("Download the Procedure Manual as PDF:")
            pdf_filename = generate_pdf(procedure_manual)
            with open(pdf_filename, "rb") as pdf_file:
                st.download_button("Download Procedure Manual", pdf_file, file_name="procedure_manual.pdf")

        except Exception as e:
            st.error(f"Error processing the documents: {e}")

if __name__ == "__main__":
    generate_procedure_manual_page()
